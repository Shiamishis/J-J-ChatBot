"""
Parses all training material files (PDF, DOCX, XLSX, MP4) under
`data/Dummy Training Materials/` into per-dashboard plain text files
under `data/training_context/`.

One subfolder = one dashboard. Output file is named after the subfolder
(slugified). The resulting files are loaded by Resources at startup and
served by the DashboardHandler.

Supported file types:
  - .pdf   -> pypdf text extraction
  - .docx  -> python-docx paragraphs + tables
  - .xlsx  -> openpyxl, all sheets dumped as pipe-delimited rows
  - .mp4   -> Whisper transcription via Swiss AI (preferred) or Groq
              (set SWISSAI_API_KEY and/or GROQ_API_KEY; files >24MB skipped)

Run directly:
    python scripts/parse_training_materials.py

Or as part of the full setup pipeline:
    python scripts/run_setup.py
"""

from __future__ import annotations

import os
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
SOURCE_DIR = ROOT / "data" / "Dummy Training Materials"
OUTPUT_DIR = ROOT / "data" / "training_context"

# Groq audio API limit is 25MB; leave a small margin.
MAX_AUDIO_BYTES = 24 * 1024 * 1024


def parse_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return f"[pypdf not installed, cannot parse {path.name}]"
    try:
        reader = PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
            except Exception as e:
                text = f"[page {i} extract_error={e}]"
            if text.strip():
                pages.append(text.strip())
        return "\n\n".join(pages)
    except Exception as e:
        return f"[pdf load_error={e}]"


def parse_docx(path: Path) -> str:
    try:
        import docx
    except ImportError:
        return f"[python-docx not installed, cannot parse {path.name}]"
    try:
        doc = docx.Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        return f"[docx load_error={e}]"


def parse_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return f"[openpyxl not installed, cannot parse {path.name}]"
    try:
        wb = load_workbook(str(path), read_only=True, data_only=True)
        out_parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            out_parts.append(f"### Sheet: {sheet_name}")
            for row in ws.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                if any(c.strip() for c in cells):
                    out_parts.append(" | ".join(cells))
        wb.close()
        return "\n".join(out_parts)
    except Exception as e:
        return f"[xlsx load_error={e}]"


def _transcribe_swissai(path: Path, api_key: str) -> str:
    """Try Swiss AI's OpenAI-compatible endpoint for audio transcription."""
    try:
        from openai import OpenAI
    except ImportError:
        return f"[openai package not installed, cannot transcribe {path.name}]"
    try:
        client = OpenAI(
            base_url="https://api.swissai.cscs.ch/v1",
            api_key=api_key,
        )
        with open(path, "rb") as f:
            transcript = client.audio.transcriptions.create(
                file=f,
                model="whisper-large-v3",
                response_format="text",
            )
        if hasattr(transcript, "text"):
            return transcript.text.strip()
        return str(transcript).strip()
    except Exception as e:
        return f"[swissai_transcription_error={e}]"


def _transcribe_groq(path: Path, api_key: str) -> str:
    """Fallback: use Groq's Whisper API."""
    try:
        from groq import Groq
    except ImportError:
        return f"[groq package not installed, cannot transcribe {path.name}]"
    try:
        client = Groq(api_key=api_key)
        with open(path, "rb") as f:
            transcript = client.audio.transcriptions.create(
                file=(path.name, f.read()),
                model="whisper-large-v3",
                response_format="text",
            )
        if hasattr(transcript, "text"):
            return transcript.text.strip()
        return str(transcript).strip()
    except Exception as e:
        return f"[groq_transcription_error={e}]"


def parse_mp4(path: Path) -> str:
    size = path.stat().st_size
    if size > MAX_AUDIO_BYTES:
        return (
            f"[{path.name} is {size / 1024 / 1024:.1f}MB, exceeds the "
            f"25MB upload limit; transcription skipped]"
        )

    swiss_key = os.environ.get("SWISSAI_API_KEY")
    groq_key = os.environ.get("GROQ_API_KEY")

    if swiss_key:
        result = _transcribe_swissai(path, swiss_key)
        # If Swiss AI failed (no Whisper model, etc.) and Groq key exists, fall back.
        if result.startswith("[swissai_transcription_error=") and groq_key:
            print(f"    Swiss AI transcription failed, falling back to Groq for {path.name}")
            return _transcribe_groq(path, groq_key)
        return result

    if groq_key:
        return _transcribe_groq(path, groq_key)

    return (
        f"[no transcription provider available for {path.name}; "
        f"set SWISSAI_API_KEY or GROQ_API_KEY]"
    )


PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".xlsx": parse_xlsx,
    ".mp4": parse_mp4,
}


def parse_file(path: Path) -> str:
    ext = path.suffix.lower()
    parser = PARSERS.get(ext)
    if parser is None:
        return f"[no parser for extension {ext}]"
    print(f"  - {path.name} ({ext})")
    return parser(path)


def parse_dashboard_folder(folder: Path) -> str:
    sections = []
    for f in sorted(folder.iterdir()):
        if not f.is_file():
            continue
        text = parse_file(f)
        sections.append(f"=== File: {f.name} ===\n{text}")
    return "\n\n".join(sections)


def slugify(name: str) -> str:
    out = []
    for ch in name:
        if ch.isalnum() or ch in ("_", "-"):
            out.append(ch)
        elif ch == " ":
            out.append("_")
    return "".join(out)


def main() -> int:
    if not SOURCE_DIR.exists():
        print(f"Source folder not found: {SOURCE_DIR}")
        return 1

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    dashboard_folders = [p for p in SOURCE_DIR.iterdir() if p.is_dir()]
    if not dashboard_folders:
        print(f"No dashboard subfolders found under {SOURCE_DIR}")
        return 1

    for folder in sorted(dashboard_folders):
        print(f"\nDashboard: {folder.name}")
        content = parse_dashboard_folder(folder)
        out_path = OUTPUT_DIR / f"{slugify(folder.name)}.txt"
        out_path.write_text(content, encoding="utf-8")
        print(f"  -> wrote {out_path.relative_to(ROOT)} ({len(content)} chars)")

    print("\nDone.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
