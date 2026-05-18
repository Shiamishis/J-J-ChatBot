#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import sqlite3
from pathlib import Path
from typing import Any

from docx import Document

ROOT = Path(__file__).resolve().parents[1]


# ---------------------------------------------------------------------------
# Documentation parsing (Robust Table Discovery)
# ---------------------------------------------------------------------------

def find_table_by_headers(doc: Document, required_headers: list[str]) -> tuple[Any, list[str]]:
    """
    Scans all tables in the document. Returns the first table that contains
    all the required_headers (case-insensitive) in its first row.
    """
    for table in doc.tables:
        if not table.rows:
            continue

        # Get text from the first row and normalize it
        header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]

        # Check if all required keywords are present in the header
        if all(any(req.lower() in h for h in header_cells) for req in required_headers):
            return table, header_cells

    req_str = ", ".join(required_headers)
    raise ValueError(f"Could not find a table containing headers: [{req_str}]")


def get_table_roles(documentation_file: Path) -> dict[str, str]:
    """
    Identifies the Inventory table and returns {table_name: role}.
    """
    doc = Document(str(documentation_file))
    # Requirement: The Inventory table must have "Table", "Type", and "Description"
    table, header = find_table_by_headers(doc, ["table", "type", "description"])

    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    col_table = header.index("table")
    col_type = header.index("type")
    col_desc = header.index("description")

    roles: dict[str, str] = {}
    for row in rows[1:]:
        name = row[col_table]
        if not name:
            continue
        if "bridge" in row[col_desc].lower():
            roles[name] = "bridge"
        elif row[col_type].lower().startswith("fact"):
            roles[name] = "fact"
        else:
            roles[name] = "dimension"
    return roles


def get_primary_keys(documentation_file: Path) -> dict[str, str]:
    """
    Identifies the Inventory table and returns {table_name: pk_column}.
    """
    doc = Document(str(documentation_file))
    # Requirement: The Inventory table must have "Table" and "Primary Key"
    table, header = find_table_by_headers(doc, ["table", "primary key"])

    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    col_table = header.index("table")
    col_pk = header.index("primary key")

    pks: dict[str, str] = {}
    for row in rows[1:]:
        name = row[col_table]
        if not name:
            continue
        pk_raw = row[col_pk]
        pk_col = pk_raw.split("+")[-1].strip() if "+" in pk_raw else pk_raw
        pks[name] = pk_col
    return pks


def get_edges(documentation_file: Path) -> list[tuple[str, str, str, str, str, float]]:
    """
    Identifies the Relationships table and returns a list of edge tuples.
    """
    doc = Document(str(documentation_file))
    # Requirement: The Relationships table must have "From Table" and "To Table"
    table, header = find_table_by_headers(doc, ["from table", "from column", "to table", "to column"])

    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    col_src_table = header.index("from table")
    col_src_col = header.index("from column")
    col_tgt_table = header.index("to table")
    col_tgt_col = header.index("to column")

    edges: list[tuple[str, str, str, str, str, float]] = []
    for row in rows[1:]:
        src_table = row[col_src_table]
        if not src_table:
            continue
        src_col = row[col_src_col]
        tgt_table = row[col_tgt_table]
        tgt_col = row[col_tgt_col]

        if src_table == tgt_table:
            edges.append((src_table, src_col, tgt_table, tgt_col, "documented_self_fk", 0.95))
        else:
            edges.append((src_table, src_col, tgt_table, tgt_col, "documented_fk", 1.0))
    return edges


# ---------------------------------------------------------------------------
# DB helpers & Metadata Logic (Unchanged from original structure)
# ---------------------------------------------------------------------------

def existing_tables(conn: sqlite3.Connection) -> set[str]:
    rows = conn.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'").fetchall()
    return {r[0] for r in rows}


def table_columns(conn: sqlite3.Connection, table: str) -> set[str]:
    rows = conn.execute(f'PRAGMA table_info("{table}")').fetchall()
    return {r[1] for r in rows}


def create_metadata_schema(conn: sqlite3.Connection) -> None:
    conn.executescript("""
    CREATE TABLE IF NOT EXISTS _metadata_tables (
        table_name TEXT PRIMARY KEY,
        table_role TEXT,
        row_count INTEGER NOT NULL DEFAULT 0,
        column_count INTEGER NOT NULL DEFAULT 0,
        primary_key_column TEXT
    );
    CREATE TABLE IF NOT EXISTS _metadata_columns (
        table_name TEXT NOT NULL,
        column_name TEXT NOT NULL,
        sqlite_type TEXT,
        is_pk INTEGER NOT NULL DEFAULT 0,
        is_not_null INTEGER NOT NULL DEFAULT 0,
        ordinal_position INTEGER NOT NULL,
        PRIMARY KEY (table_name, column_name)
    );
    CREATE TABLE IF NOT EXISTS _metadata_foreign_keys (
        source_table TEXT NOT NULL,
        source_column TEXT NOT NULL,
        target_table TEXT NOT NULL,
        target_column TEXT NOT NULL,
        relation_type TEXT NOT NULL,
        confidence REAL NOT NULL,
        PRIMARY KEY (source_table, source_column, target_table, target_column)
    );
    CREATE TABLE IF NOT EXISTS _metadata_semantic_columns (
        table_name TEXT NOT NULL,
        column_name TEXT NOT NULL,
        semantic_role TEXT NOT NULL,
        PRIMARY KEY (table_name, column_name, semantic_role)
    );
    """)


def semantic_role(column: str) -> str:
    c = column.lower()
    if c.endswith("_id") or c == "id": return "id"
    if "time" in c or c in {"year", "quarter", "month", "fiscal_year"}: return "time"
    if c.startswith("is_") or c.endswith("_flag"): return "flag"
    if any(m in c for m in ["revenue", "units", "count", "value", "score", "volume", "pct"]): return "metric"
    return "dimension_attr"


# def rebuild_metadata(conn: sqlite3.Connection, table_roles: dict, primary_keys: dict, edges: list) -> dict:
#     create_metadata_schema(conn)
#     conn.execute("DELETE FROM _metadata_tables")
#     conn.execute("DELETE FROM _metadata_columns")
#     conn.execute("DELETE FROM _metadata_foreign_keys")
#     conn.execute("DELETE FROM _metadata_semantic_columns")
#
#     tables_in_db = existing_tables(conn)
#     managed_tables = sorted([t for t in table_roles if t in tables_in_db])
#
#     for table in managed_tables:
#         cols_info = conn.execute(f'PRAGMA table_info("{table}")').fetchall()
#         row_count = conn.execute(f'SELECT COUNT(*) FROM "{table}"').fetchone()[0]
#         pk_col = primary_keys.get(table)
#
#         conn.execute("INSERT INTO _metadata_tables VALUES (?, ?, ?, ?, ?)",
#                      (table, table_roles.get(table), row_count, len(cols_info), pk_col))
#
#         for col in cols_info:
#             cid, name, col_type, notnull, _, _ = col
#             is_pk = 1 if name == pk_col else 0
#             conn.execute("INSERT INTO _metadata_columns VALUES (?, ?, ?, ?, ?, ?)",
#                          (table, name, col_type or "", is_pk, int(bool(notnull)), cid))
#             conn.execute("INSERT INTO _metadata_semantic_columns VALUES (?, ?, ?)", (table, name, semantic_role(name)))
#
#     inserted_edges = 0
#     skipped_edges = []
#     for s_table, s_col, t_table, t_col, rel_type, conf in edges:
#         if s_table not in tables_in_db or t_table not in tables_in_db:
#             skipped_edges.append((s_table, s_col, t_table, t_col, "table_missing"))
#             continue
#         if s_col not in table_columns(conn, s_table) or t_col not in table_columns(conn, t_table):
#             skipped_edges.append((s_table, s_col, t_table, t_col, "column_missing"))
#             continue
#
#         conn.execute("INSERT OR REPLACE INTO _metadata_foreign_keys VALUES (?, ?, ?, ?, ?, ?)",
#                      (s_table, s_col, t_table, t_col, rel_type, conf))
#         inserted_edges += 1
#
#     conn.commit()
#     return {"managed_tables": managed_tables, "inserted_edges": inserted_edges, "skipped_edges": skipped_edges}

import networkx as nx
import warnings


def rebuild_metadata(conn: sqlite3.Connection, table_roles: dict, primary_keys: dict, edges: list) -> dict:
    create_metadata_schema(conn)
    # ... (standard DELETEs) ...
    conn.execute("DELETE FROM _metadata_tables")
    conn.execute("DELETE FROM _metadata_columns")
    conn.execute("DELETE FROM _metadata_foreign_keys")
    conn.execute("DELETE FROM _metadata_semantic_columns")

    tables_in_db = existing_tables(conn)

    # 1. First, identify ALL potential valid edges
    valid_edges = []
    active_tables_in_edges = set()

    for s_table, s_col, t_table, t_col, rel_type, conf in edges:
        if s_table in tables_in_db and t_table in tables_in_db:
            s_cols = table_columns(conn, s_table)
            t_cols = table_columns(conn, t_table)
            if s_col in s_cols and t_col in t_cols:
                valid_edges.append({
                    "source_table": s_table, "source_column": s_col,
                    "target_table": t_table, "target_column": t_col,
                    "relation_type": rel_type, "confidence": conf
                })
                active_tables_in_edges.add(s_table)
                active_tables_in_edges.add(t_table)

    # 2. Generalizable Filtering: Only manage tables that have at least one edge
    # This automatically excludes FACT_KPI_Summary or any other "orphans"
    managed_tables = sorted([t for t in table_roles if t in tables_in_db and t in active_tables_in_edges])

    # 3. Connectivity Analysis (The "Island" Warning)
    G = nx.Graph()
    G.add_nodes_from(managed_tables)
    for e in valid_edges:
        G.add_edge(e["source_table"], e["target_table"])

    components = list(nx.connected_components(G))
    large_clusters = [c for c in components if len(c) > 1]

    if len(large_clusters) > 1:
        print("\n" + "!" * 60)
        print("WARNING: DISCONNECTED SCHEMA DETECTED")
        print(f"Found {len(large_clusters)} separate clusters of connected tables.")
        for i, cluster in enumerate(large_clusters, 1):
            print(f"  Cluster {i}: {cluster}")
        print("This may cause the LLM to fail when joining across these clusters.")
        print("!" * 60 + "\n")

    # 4. Proceed with DB insertion for managed tables only
    for table in managed_tables:
        cols_info = conn.execute(f'PRAGMA table_info("{table}")').fetchall()
        row_count = conn.execute(f'SELECT COUNT(*) FROM "{table}"').fetchone()[0]
        pk_col = primary_keys.get(table)

        conn.execute("INSERT INTO _metadata_tables VALUES (?, ?, ?, ?, ?)",
                     (table, table_roles.get(table), row_count, len(cols_info), pk_col))

        for col in cols_info:
            cid, name, col_type, notnull, _, _ = col
            is_pk = 1 if name == pk_col else 0
            conn.execute("INSERT INTO _metadata_columns VALUES (?, ?, ?, ?, ?, ?)",
                         (table, name, col_type or "", is_pk, int(bool(notnull)), cid))
            conn.execute("INSERT INTO _metadata_semantic_columns VALUES (?, ?, ?)",
                         (table, name, semantic_role(name)))

    # 5. Insert valid edges into DB
    for e in valid_edges:
        conn.execute("INSERT OR REPLACE INTO _metadata_foreign_keys VALUES (?, ?, ?, ?, ?, ?)",
                     (e["source_table"], e["source_column"], e["target_table"], e["target_column"],
                      e["relation_type"], e["confidence"]))

    conn.commit()

    return {
        "managed_tables": managed_tables,
        "valid_edges": valid_edges,
        "skipped_edges": [e for e in edges if e[0] not in active_tables_in_edges and e[2] not in active_tables_in_edges]
    }

def export_json(conn: sqlite3.Connection, path: Path) -> None:
    nodes = [r[0] for r in conn.execute("SELECT table_name FROM _metadata_tables ORDER BY table_name").fetchall()]
    edges = conn.execute(
        "SELECT source_table, source_column, target_table, target_column, relation_type, confidence FROM _metadata_foreign_keys ORDER BY source_table, source_column, target_table").fetchall()

    payload = {
        "nodes": nodes,
        "edges": [{"source_table": r[0], "source_column": r[1], "target_table": r[2], "target_column": r[3],
                   "relation_type": r[4], "confidence": r[5]} for r in edges],
    }
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")

def export_documentation_txt(documentation_file: Path) -> Path:
    """
    Converts the technical documentation .docx to plain text and saves it
    under the data folder alongside the source file.
    """
    doc = Document(str(documentation_file))
    lines = []

    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

        if tag == "p":
            from docx.text.paragraph import Paragraph
            text = Paragraph(block, doc).text.strip()
            if text:
                lines.append(text)

        elif tag == "tbl":
            from docx.table import Table
            for row in Table(block, doc).rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                lines.append("\t".join(cells))
            lines.append("")  # blank line after each table

    out_path = documentation_file.parent / "documentation.txt"
    out_path.write_text("\n".join(lines), encoding="utf-8")
    return out_path

def main() -> None:
    parser = argparse.ArgumentParser(description="Automated metadata bootstrap for NovaCarta SQLite DB.")
    parser.add_argument("--db", default=str(ROOT / "local.db"), help="Path to SQLite DB")
    parser.add_argument("--out", default=str(ROOT / "data" / "metadata_graph_manual.json"), help="Output graph JSON")
    args = parser.parse_args()

    db_path = Path(args.db).resolve()
    data_folder_path = ROOT / "data"

    # Auto-find documentation file
    documentation_file = next(
        (f for f in data_folder_path.iterdir() if f.suffix == ".docx" and "Technical_Documentation" in f.name), None)
    if not documentation_file:
        raise FileNotFoundError(f"Technical_Documentation .docx not found in {data_folder_path}")

    # Perform Extraction using new robust methods
    table_roles = get_table_roles(documentation_file)
    primary_keys = get_primary_keys(documentation_file)
    edges = get_edges(documentation_file)

    with sqlite3.connect(str(db_path)) as conn:
        summary = rebuild_metadata(conn, table_roles, primary_keys, edges)
        export_json(conn, Path(args.out).resolve())

    export_documentation_txt(documentation_file)

    print(f"Success! Metadata built for {len(summary['managed_tables'])} tables.")
    print(f"Edges: {summary['valid_edges']} inserted, {len(summary['skipped_edges'])} skipped.")


if __name__ == "__main__":
    main()
